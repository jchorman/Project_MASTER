import os
import sys
import shutil
import re
import win32com.client
import win32com
from docx import Document
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date, datetime
import configparser
import platform

# Determine the paths based on execution context
if getattr(sys, 'frozen', False):
    config_path = os.path.join(sys._MEIPASS, 'config.ini')
    email_body_path = os.path.join(sys._MEIPASS, 'email_body.html')
else:
    config_path = 'config.ini'
    email_body_path = 'email_body.html'

# Load and read the configuration file
config = configparser.ConfigParser()
config.read(config_path)

# Adjust the paths in the config values
BASE_DIRECTORY = os.path.join(config['Paths']['base_directory'])
WORKBOOK_PATH = os.path.join(config['Paths']['workbook_path'])
EMAIL_TEMPLATE = os.path.join(config['Paths']['email_template'])
EMAIL_BODY = os.path.join(config['Paths']['email_body'])
SHARED_MAILBOX_EMAIL = config['Settings']['shared_mailbox_email']
ON_PREM_VISIO_TEMPLATE = os.path.join(config['Paths']['on_prem_visio_template'])
OFF_PREM_VISIO_TEMPLATE = os.path.join(config['Paths']['off_prem_visio_template'])


def create_folder_and_docx(folder_path, docx_title, lsar_date=None, is_lsar=False):
    try:
        # Create the folder
        os.makedirs(folder_path, exist_ok=True)

        # Create a docx file with the specified title
        docx_file_path = os.path.join(folder_path, f'{docx_title}.docx')
        doc = Document()

        # Set the default font to Arial 12 for the entire document
        for style in doc.styles:
            if style.name == 'Normal':
                style.font.name = 'Arial'
                style.font.size = Pt(12)

        # Create a paragraph for the default text
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Set paragraph alignment to left

        if is_lsar:  # If project is starting from LSAR
            if lsar_date:  # If LSAR date is provided
                # Convert lsar_date to string before appending
                run = p.add_run(lsar_date.strftime("%m/%d/%Y"))
                run.bold = True
                p.add_run(" - LSAR held; welcome packet sent.")
        else:  # For Validated Design
            today = date.today().strftime("%m/%d/%Y")
            run = p.add_run(today)
            run.bold = True
            p.add_run(" - Project entered into system at the design phase.")

        doc.save(docx_file_path)

        return docx_file_path
    except Exception as e:
        print(f"An error occurred: {str(e)}")


def download_attachments_from_calendar(shared_calendar, attachment_dir, lsar_date):
    appts = shared_calendar.Items
    appts.Sort("[Start]")
    appts.IncludeRecurrences = "True"

    filter_str = "[Start] >= '" + lsar_date.strftime("%m/%d/%Y") + " 08:00 AM' AND [End] <= '" + lsar_date.strftime("%m/%d/%Y") + " 04:00 PM'"
    selected_appts = appts.Restrict(filter_str)

    image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tif', '.tiff']

    if not selected_appts:
        print(f"No meetings found for {lsar_date}.")
    else:
        os.makedirs(attachment_dir, exist_ok=True)
        attachments_found = False

        for a in selected_appts:
            if "lsar" in a.Subject.strip().lower():
                for attachment in a.Attachments:
                    if attachment.Type == 1 and not any(attachment.FileName.lower().endswith(ext) for ext in image_extensions):
                        attachment_filename = os.path.join(attachment_dir, attachment.FileName)
                        try:
                            attachment.SaveAsFile(attachment_filename)
                            print(f"Downloaded attachment: {attachment_filename}")
                            attachments_found = True
                        except Exception as e:
                            print(f"Error saving attachment: {e}")

        if attachments_found:
            print(f"Attachments downloaded to the directory: {attachment_dir}")
        else:
            print(f"No attachments found for meetings on {lsar_date}.")


def get_email_body_from_template(agency_name, project_name, original_project_name):
    global email_body_path  # This is to clarify that we're using the global variable, though not strictly needed here
    with open(email_body_path, 'r') as file:
        email_content = file.read()  # Use a different variable name here

    return email_content.format(agency_name=agency_name, project_name=original_project_name)


def search_and_select_project(base_directory):
    try:
        # Prompt the user to enter a keyword
        keyword = input(
            "Enter a keyword to search for a project: ").lower()  # Convert the keyword to lowercase for case-insensitive search

        # Search for project folders containing the keyword
        matching_projects = []
        for root, dirs, _ in os.walk(os.path.join(base_directory, "Active Projects")):
            for dir_name in dirs:
                if keyword in dir_name.lower():  # Convert folder name to lowercase for case-insensitive comparison
                    matching_projects.append(dir_name)

        if not matching_projects:
            print(f"No projects found matching the keyword '{keyword}'.")
            return None

        # Display search results with numbers for selection
        print("Search results:")
        for i, project_name in enumerate(matching_projects, start=1):
            print(f"{i}. {project_name}")

        # Prompt user to select a project by number
        while True:
            try:
                selection = int(input("Enter the number of the project you want to close: "))
                if 1 <= selection <= len(matching_projects):
                    selected_project_name = matching_projects[selection - 1]
                    return os.path.join(base_directory, "Active Projects", selected_project_name)
                else:
                    print("Invalid selection. Please enter a valid number.")
            except ValueError:
                print("Invalid input. Please enter a valid number.")

    except Exception as e:
        print(f"An error occurred during project selection: {str(e)}")


def format_project_name(project_name):
    """Convert each word in the project name to start with an uppercase letter and join them together."""
    words = project_name.split()
    capitalized_words = [word.capitalize() for word in words]
    return ''.join(capitalized_words)


def create_projects(workbook, num_projects):
    base_directory = BASE_DIRECTORY

    # Select the active sheet (you may need to specify the sheet name if it's not the active one)
    sheet = workbook['Yes']

    try:
        for i in range(num_projects):
            # Read column headers from the first row of the spreadsheet
            headers = [cell.value for cell in sheet[1]]
            data = []
            project_type_map = {
                "1": "On-Prem",
                "2": "Off-Prem"
            }

            for i, header in enumerate(headers):
                while True:  # This loop will continue until a valid input is received
                    # Change the prompt if the header is "On-Prem or Off-Prem"
                    prompt_message = f"On-Prem (Enter '1') or Off-Prem (Enter '2')" if header == "On-Prem or Off-Prem" else header
                    value = input(f"{prompt_message}: ").strip()

                    if header == "On-Prem or Off-Prem":
                        # Check if the entered value is in our map (i.e., "1" or "2")
                        if value in project_type_map:
                            # Convert the value to "On-Prem" or "Off-Prem"
                            value = project_type_map[value]
                            break
                        else:
                            print("Invalid choice. Enter '1' for On-Prem or '2' for Off-Prem.")
                        continue
                    else:
                        break

                data.append(value)

            # Get the LSAR date from the data entered by the user (assuming it's the first column)
            lsar_date = data[0]  # LSAR date is in the first column
            lsar_date = datetime.strptime(lsar_date, "%m/%d/%Y")

            # Combine agency name and project name with a period
            agency_name = data[2]
            division_name = data[3]
            original_project_name = data[4]
            project_name = format_project_name(original_project_name)
            project_type = data[5].strip().lower()

            # Combine the names as per the given format
            if division_name and division_name.strip() != "":
                # If there is a division name, format the project name with it
                formatted_project_name = f"{agency_name}-{division_name}.{project_name.replace(' ', '')}"
            else:
                # If there is no division name, use the old format
                formatted_project_name = f"{agency_name}.{project_name.replace(' ', '')}"

            docx_title = f"{formatted_project_name}.STATUS_SHEET"

            # Create a folder with the specified name
            folder_path = os.path.join(base_directory, "Active Projects", formatted_project_name)
            docx_file_path = create_folder_and_docx(folder_path, docx_title, lsar_date, is_lsar=True)

            # Create a folder for attachments within the project folder
            attachment_dir = os.path.join(folder_path, "LSAR Meeting Documents")
            os.makedirs(attachment_dir, exist_ok=True)  # Ensure the directory exists

            # Download the attachments from the calendar
            download_attachments_from_calendar(shared_calendar, attachment_dir, lsar_date)

            if docx_file_path:
                try:
                    # Define Visio template paths based on the extracted project_type
                    if project_type == "on-prem":
                        visio_template_path = ON_PREM_VISIO_TEMPLATE
                        new_visio_title = f"OnPrem-{formatted_project_name}-DESIGN"
                    elif project_type == "off-prem":
                        visio_template_path = OFF_PREM_VISIO_TEMPLATE
                        new_visio_title = f"OffPrem-{formatted_project_name}-DESIGN"

                    # Copy the selected Visio template to the project folder
                    if os.path.exists(visio_template_path):
                        shutil.copy(visio_template_path, folder_path)
                    else:
                        print(f"Error: Visio template not found at '{visio_template_path}'")

                    # Rename the copied Visio template to the new title
                    copied_visio_path = os.path.join(folder_path, "OnPrem-SA-Example-Diagram.vsdx") if project_type == "on-prem" else os.path.join(folder_path, "OffPrem-SA-Example-Diagram.vsdx")
                    new_copied_visio_path = os.path.join(folder_path, f'{new_visio_title}.vsdx')
                    os.rename(copied_visio_path, new_copied_visio_path)

                    print(f"Project folder created at '{folder_path}'")
                    print(f"Docx file created at '{docx_file_path}'")
                    print(f"Visio template renamed to '{new_visio_title}.vsdx'")

                    # Get the email template
                    oft_file_path = EMAIL_TEMPLATE

                    # Create a new mail item from the template
                    mail_item = outlook.CreateItemFromTemplate(oft_file_path)

                    # Append the original project name to the email subject with spaces and replace the period with " - "
                    project_parts = formatted_project_name.split(".")
                    if len(project_parts) == 2:
                        left_part, _ = project_parts  # Use the formatted agency and division but use the original project name for the right part
                        subject_with_spaces = f"{left_part} - {original_project_name}"
                        mail_item.Subject = f"{mail_item.Subject} - {subject_with_spaces}"
                    else:
                        print("Invalid project name format. It should contain exactly one period.")

                    # Construct the HTML email body with placeholders
                    email_body = get_email_body_from_template(agency_name, project_name, original_project_name)

                    # Set the email body with HTML formatting
                    mail_item.HTMLBody = email_body

                    # Attach the Visio file to the email
                    visio_file_path = os.path.join(folder_path, f'{new_visio_title}.vsdx')
                    mail_item.Attachments.Add(visio_file_path)

                    # Save the email as a .msg file in the project folder
                    email_msg_file_path = os.path.join(folder_path, f'{project_name}_email.msg')

                    # Using 3 as the integer value for olMSG to save the email as .msg format
                    mail_item.SaveAs(email_msg_file_path, 3)

                    # Enter data into the appropriate columns
                    sheet.append([*data])

                    # Highlight the entire row as yellow
                    yellow_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
                    for cell in sheet[sheet.max_row]:
                        cell.font = Font(name='Arial', size=11)
                        cell.alignment = Alignment(horizontal='right') if cell.column in [1, 7] else Alignment(horizontal='left')
                        cell.fill = yellow_fill

                    # Save the updated Excel file
                    workbook.save(WORKBOOK_PATH)

                    print("Email template created, Visio attached, data entered, and row highlighted successfully.")
                    print(f".msg file copied to project folder: '{email_msg_file_path}'")

                except Exception as e:
                    print(f"An error occurred: {str(e)}")

    except Exception as e:
        print(f"An error occurred during initialization: {str(e)}")


def create_projects_no_lsar(workbook, num_projects):
    base_directory = BASE_DIRECTORY
    sheet = workbook['Yes']

    try:
        for i in range(num_projects):
            headers = [cell.value for cell in sheet[1]]
            data = []
            project_type_map = {
                "1": "On-Prem",
                "2": "Off-Prem"
            }

            for i, header in enumerate(headers):
                # Automatically set default values for columns 0 and 1
                if i == 0:  # LSAR Date
                    data.append('')  # Blank value
                    continue
                elif i == 1:  # eReview
                    data.append('NA - Design')
                    continue

                while True:
                    prompt_message = f"On-Prem (Enter '1') or Off-Prem (Enter '2')" if header == "On-Prem or Off-Prem" else header
                    value = input(f"{prompt_message}: ").strip()

                    if header == "On-Prem or Off-Prem":
                        if value in project_type_map:
                            value = project_type_map[value]
                            break
                        else:
                            print("Invalid choice. Enter '1' for On-Prem or '2' for Off-Prem.")
                    else:
                        break

                data.append(value)

            current_date = datetime.now().strftime("%m/%d/%Y")
            agency_name = data[2]
            division_name = data[3]
            original_project_name = data[4]
            project_name = format_project_name(original_project_name)
            project_type = data[5].strip().lower()

            if division_name and division_name.strip() != "":
                formatted_project_name = f"{agency_name}-{division_name}.{project_name.replace(' ', '')}"
            else:
                formatted_project_name = f"{agency_name}.{project_name.replace(' ', '')}"

            docx_title = f"{formatted_project_name}.STATUS_SHEET"
            folder_path = os.path.join(base_directory, "Active Projects", formatted_project_name)
            docx_file_path = create_folder_and_docx(folder_path, docx_title, is_lsar=False)

            if project_type == "on-prem":
                visio_template_path = ON_PREM_VISIO_TEMPLATE
                new_visio_title = f"OnPrem-{formatted_project_name}-DESIGN"
            elif project_type == "off-prem":
                visio_template_path = OFF_PREM_VISIO_TEMPLATE
                new_visio_title = f"OffPrem-{formatted_project_name}-DESIGN"

            if os.path.exists(visio_template_path):
                shutil.copy(visio_template_path, folder_path)
            else:
                print(f"Error: Visio template not found at '{visio_template_path}'")

            copied_visio_path = os.path.join(folder_path, "OnPrem-SA-Example-Diagram.vsdx") if project_type == "on-prem" else os.path.join(folder_path, "OffPrem-SA-Example-Diagram.vsdx")
            new_copied_visio_path = os.path.join(folder_path, f'{new_visio_title}.vsdx')
            os.rename(copied_visio_path, new_copied_visio_path)

            print(f"Project folder created at '{folder_path}'")
            print(f"Docx file created at '{docx_file_path}'")
            print(f"Visio template renamed to '{new_visio_title}.vsdx'")

            oft_file_path = EMAIL_TEMPLATE
            mail_item = outlook.CreateItemFromTemplate(oft_file_path)
            project_parts = formatted_project_name.split(".")
            if len(project_parts) == 2:
                left_part, _ = project_parts
                subject_with_spaces = f"{left_part} - {original_project_name}"
                mail_item.Subject = f"{mail_item.Subject} - {subject_with_spaces}"
            else:
                print("Invalid project name format. It should contain exactly one period.")

            email_body = get_email_body_from_template(agency_name, project_name, original_project_name)
            mail_item.HTMLBody = email_body
            visio_file_path = os.path.join(folder_path, f'{new_visio_title}.vsdx')
            mail_item.Attachments.Add(visio_file_path)
            email_msg_file_path = os.path.join(folder_path, f'{project_name}_email.msg')
            mail_item.SaveAs(email_msg_file_path, 3)

            sheet.append([*data])
            yellow_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
            for cell in sheet[sheet.max_row]:
                cell.font = Font(name='Arial', size=11)
                cell.alignment = Alignment(horizontal='right') if cell.column in [1, 7] else Alignment(horizontal='left')
                cell.fill = yellow_fill

            workbook.save(WORKBOOK_PATH)
            print("Email template created, Visio attached, data entered, and row highlighted successfully.")
            print(f".msg file copied to project folder: '{email_msg_file_path}'")

    except Exception as e:
        print(f"An error occurred: {str(e)}")


def search_directory(keyword, directory_path):
    matched_folders = []

    print(f"Checking directory {directory_path}...")  # Diagnostic
    for foldername, subfolders, filenames in os.walk(directory_path):
        match_found = False  # Flag to track if a match is found in the current folder

        # Check folder name for a match
        if keyword.lower() in foldername.lower():
            matched_folders.append(foldername)
            match_found = True

        # If no match found in folder name, check filenames
        if not match_found:
            for filename in filenames:
                if keyword.lower() in filename.lower():
                    matched_folders.append(foldername)
                    break  # Exit the loop once a match is found

    if not matched_folders:
        print(f"No matches found in directory {directory_path} for keyword '{keyword}'")  # Diagnostic            
    return matched_folders


def search_spreadsheet(keyword, spreadsheet_path):
    workbook = openpyxl.load_workbook(spreadsheet_path)

    matched_rows = []

    print(f"Checking workbook {spreadsheet_path}...")  # Diagnostic
    for sheet in workbook.worksheets:
        for row_idx, row in enumerate(sheet.iter_rows(), start=1):  # Using start=1 for 1-based row numbering
            for cell in row:
                if cell.value and keyword.lower() in str(cell.value).lower():
                    matched_rows.append((sheet.title, row_idx))
                    break  # Break out of the cell loop as we found a match in this row

    if not matched_rows:
        print(f"No matches found in workbook {spreadsheet_path} for keyword '{keyword}'")  # Diagnostic
    return matched_rows


def get_directories_matching_keyword(keyword, base_directory):
    """Return a list of directories that match the given keyword, one level below the base_directory."""
    matching_dirs = []

    # Get all entries in the base_directory
    all_entries = os.listdir(base_directory)
    for entry in all_entries:
        full_path = os.path.join(base_directory, entry)
        # Check if the entry is a directory and matches the keyword
        if os.path.isdir(full_path) and keyword.lower() in entry.lower():
            matching_dirs.append(entry)  # Append only the directory name, not the full path

    return matching_dirs


def search_for_project(base_directory):
    print("Please specify your search option:")
    print("1. Keyword search across the entire base directory.")
    print("2. Keyword search within a specific directory.")
    print("3. List all directories.")

    choice = input("Choice: ").strip()

    # If user selects '1' or starts typing, consider it a search across the entire directory
    if choice == '1' or not choice.isnumeric():
        keyword = choice if not choice.isnumeric() else input("Enter the keyword for search: ")
        directory_path = BASE_DIRECTORY
    elif choice == '2':
        directory_name = input("Enter the specific directory name (or a part of it): ").strip()
        directory_path = os.path.join(BASE_DIRECTORY, directory_name)

        # If the directory doesn't exist or is only a partial match
        if not os.path.isdir(directory_path) or directory_name == "":
            # Get a list of directories that contain the entered string
            matching_directories = get_directories_matching_keyword(directory_name, BASE_DIRECTORY)

            # If no matching directories found
            if not matching_directories:
                print(f"No directories containing '{directory_name}' found. Defaulting to base directory.")
                directory_path = BASE_DIRECTORY
            else:
                # Print matching directories for selection
                for idx, directory in enumerate(matching_directories, 1):
                    print(f"{idx}. {directory}")

                # Ask the user to select from the list of matching directories
                while True:
                    try:
                        dir_selection = int(input("Enter the number of the directory you want to search: "))
                        if 1 <= dir_selection <= len(matching_directories):
                            directory_path = os.path.join(BASE_DIRECTORY, matching_directories[dir_selection - 1])
                            break
                        else:
                            print("Invalid selection. Please enter a number from the list.")
                    except ValueError:
                        print("Please enter a valid number.")

        keyword = input("Enter the keyword for search: ")
    elif choice == '3':
        all_directories = [entry for entry in os.listdir(base_directory) if os.path.isdir(os.path.join(base_directory, entry))]
        for idx, directory in enumerate(all_directories, 1):
            print(f"{idx}. {directory}")

        # Ask the user to select from the list of all directories
        while True:
            try:
                dir_selection = int(input("Enter the number of the directory you want to search (or '0' to go back): "))
                if 0 <= dir_selection <= len(all_directories):
                    if dir_selection == 0:  # Go back to main menu or quit
                        return  # or break, depending on the rest of your logic
                    directory_path = os.path.join(BASE_DIRECTORY, all_directories[dir_selection - 1])
                    break
                else:
                    print("Invalid selection. Please enter a number from the list.")
            except ValueError:
                print("Please enter a valid number.")

        keyword = input("Enter the keyword for search (or press 'Enter' to skip search): ").strip()
        if not keyword:  # If the user doesn't provide a keyword, return to the main menu or quit
            return  # or break, depending on the rest of your logic

    else:
        # Use the provided choice to get matching directories
        matching_directories = get_directories_matching_keyword(choice, BASE_DIRECTORY)

        # List the matching directories for selection
        for idx, directory in enumerate(matching_directories, 1):
            print(f"{idx}. {directory}")

        # Ask the user to select from the list of matching directories
        while True:
            try:
                dir_selection = int(input("Enter the number of the directory you want to search: "))
                if 1 <= dir_selection <= len(matching_directories):
                    directory_path = os.path.join(BASE_DIRECTORY, matching_directories[dir_selection - 1])
                    break
                else:
                    print("Invalid selection. Please enter a number from the list.")
            except ValueError:
                print("Please enter a valid number.")

    # Path to your spreadsheet
    spreadsheet_path = WORKBOOK_PATH
    matches_in_spreadsheet = search_spreadsheet(keyword, spreadsheet_path)

    print("\nFound in spreadsheet:")
    for sheet_name, row_num in matches_in_spreadsheet:
        print(f"Sheet: {sheet_name}, Row: {row_num}")

    # Search the specified directory (or the default BASE_DIRECTORY)
    matched_folders = search_directory(keyword, directory_path)

    while True:  # Keep prompting for folder selection
        print("\nFound in project folders:")
        for idx, folder in enumerate(matched_folders, 1):
            print(f"{idx}. {folder}")

        # Ask the user to select a folder
        while True:
            try:
                selection = int(input("\nEnter the number of the project you want to open (or '0' to exit): "))
                if 0 <= selection <= len(matched_folders):
                    break
                else:
                    print("Invalid selection. Please enter a number from the list.")
            except ValueError:
                print("Please enter a valid number.")

        # Exit the loop if the user selects '0'
        if selection == 0:
            break

        # Open the selected folder
        folder_to_open = os.path.abspath(matched_folders[selection - 1])
        if platform.system() == "Darwin":  # macOS
            os.system(f"open \"{folder_to_open}\"")
        elif platform.system() == "Windows":
            os.system(f"explorer \"{folder_to_open}\"")
        elif platform.system() == "Linux":
            os.system(f"xdg-open \"{folder_to_open}\"")

        # Optional: Clear the console screen for better user experience
        os.system('cls' if platform.system() == 'Windows' else 'clear')


def add_spaces_before_capitals(s):
    """Inserts spaces before capitals in the given string."""
    words = re.sub(r'(?<=[a-z])(?=[A-Z])', ' ', s).split()
    lowercase_words = ["and", "or", "of", "in", "to", "the", "for", "with", "on", "by", "at", "an", "as", "a", "but"]
    final_words = [word if word.lower() not in lowercase_words else word.lower() for word in words]
    return ' '.join(final_words)


def close_project_and_copy_to_validated(base_directory, project_folder):
    try:
        # Define destination paths
        closed_projects_folder = os.path.join(base_directory, "Closed Projects")
        validated_designs_folder = os.path.join(base_directory, "Validated Designs", os.path.basename(project_folder))

        # Move the project folder to "Closed Projects"
        shutil.move(project_folder, os.path.join(closed_projects_folder, os.path.basename(project_folder)))

        print(f"Project folder moved to '{closed_projects_folder}'")

        # Create a folder in "Validated Designs"
        os.makedirs(validated_designs_folder, exist_ok=True)
        print(f"Folder created in 'Validated Designs': '{validated_designs_folder}'")

        # Find the most recent .vsdx file in the project folder
        vsdx_files = [f for f in os.listdir(os.path.join(closed_projects_folder, os.path.basename(project_folder))) if f.endswith(".vsdx")]
        if vsdx_files:
            most_recent_vsdx = max(vsdx_files, key=lambda f: os.path.getmtime(os.path.join(closed_projects_folder, os.path.basename(project_folder), f)))
            most_recent_vsdx_path = os.path.join(closed_projects_folder, os.path.basename(project_folder), most_recent_vsdx)
            shutil.copy(most_recent_vsdx_path, os.path.join(validated_designs_folder, most_recent_vsdx))
            print(f"Most recent .vsdx file copied to 'Validated Designs'")
        else:
            print("No .vsdx files found in the project folder.")

        # Open the "Yes" workbook
        workbook = openpyxl.load_workbook(WORKBOOK_PATH)

        # Select the "Yes" worksheet (or specify the actual name)
        sheet = workbook['Yes']

        # Extract the ProjectName from the full project folder name
        folder_basename = os.path.basename(project_folder)

        folder_parts = folder_basename.rsplit('.', 1)
        if len(folder_parts) > 1:
            project_name_raw = folder_parts[-1]
        else:
            project_name_raw = folder_parts[0]

        # Convert ProjectName into Project Name format
        refined_project_name = add_spaces_before_capitals(project_name_raw)

        # Read the header row to find the "Project" column index
        header_row = list(sheet.iter_rows(min_row=1, max_row=1, values_only=True))[0]
        project_column_index = header_row.index("Project") + 1  # Adjust this if the column name is different   

        found_row_number = None  # Keep track of the row number where the agency was found

        # Iterate through rows to find the refined project name
        for row_number, row in enumerate(
                sheet.iter_rows(min_col=project_column_index, max_col=project_column_index, values_only=True), start=1):
            if row[0] == refined_project_name:
                found_row_number = row_number  # Store the row number where the project name was found
                break

        if found_row_number is not None:
            # Find the entire row with the project name using the stored row number
            project_row = list(sheet.iter_rows(min_row=found_row_number, max_row=found_row_number, values_only=True))[0]

            # Open the "Done" worksheet within the same workbook
            done_sheet = workbook['Done']

            # Find the next available row in the "Done" worksheet
            done_row = done_sheet.max_row + 1

            # Copy data from the project row to the "Done" worksheet
            for col, cell_value in enumerate(project_row, start=1):
                done_sheet.cell(row=done_row, column=col, value=cell_value)

            # Apply formatting to the copied row in the "Done" worksheet
            yellow_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
            for cell in done_sheet[done_row]:
                cell.font = Font(name='Arial', size=11)
                cell.alignment = Alignment(horizontal='right') if cell.column in [1, 7] else Alignment(horizontal='left')
                cell.fill = yellow_fill

            # Delete the row from the "Yes" worksheet
            sheet.delete_rows(found_row_number)

            # Save changes to the "Yes" workbook (assuming you want to save changes there)
            workbook.save(WORKBOOK_PATH)

            # Close the workbook
            workbook.close()

            print(
                f"Project entry moved to 'Done' worksheet in the 'Status_201705-OnwardCOPY' workbook, and row deleted from 'Yes'.")
        else:
            print(f"Project '{refined_project_name}' not found in the 'Yes' worksheet.")

    except Exception as e:
        print(f"An error occurred: {str(e)}")


def main():
    print("Choose an option:")
    print("1: Create a project")
    print("2: Search for a project")
    print("3: Close out a project")
    choice = input("Enter your choice (1/2/3): ")

    if choice == "1":
        print("Where will this project start (Select the number that applies)?")
        print("1) LSAR")
        print("2) Validated Design")
        start_choice = input("Enter your choice (1/2): ")

        if start_choice == "1":
            # Existing code for LSAR
            workbook = openpyxl.load_workbook(WORKBOOK_PATH)
            num_projects = int(input("Enter the number of projects to create: "))
            create_projects(workbook, num_projects)
            workbook.save(WORKBOOK_PATH)

        elif start_choice == "2":
            # Code for Validated Design
            # Omitted the part for LSAR date and LSAR related function calls
            workbook = openpyxl.load_workbook(WORKBOOK_PATH)
            num_projects = int(input("Enter the number of projects to create without LSAR: "))
            create_projects_no_lsar(workbook, num_projects)  # This function should be defined to create projects without LSAR
            workbook.save(WORKBOOK_PATH)

        else:
            print("Invalid choice for project start. Please enter either 1 or 2.")

    elif choice == "2":
        base_directory = BASE_DIRECTORY
        search_for_project(base_directory)

    elif choice == "3":
        base_directory = BASE_DIRECTORY
        project_folder = search_and_select_project(base_directory)
        if project_folder:
            close_project_and_copy_to_validated(base_directory, project_folder)

    else:
        print("Invalid choice. Please enter either 1, 2, or 3.")


if __name__ == "__main__":
    # Initialize the Outlook object once
    outlook = win32com.client.Dispatch('Outlook.Application')
    namespace = outlook.GetNamespace("MAPI")
    recipient = namespace.CreateRecipient(SHARED_MAILBOX_EMAIL)
    shared_calendar = namespace.GetSharedDefaultFolder(recipient, 9)

    while True:
        main()
        choice = input("Run script again? (y/n): ").lower()
        if choice != 'y':
            break